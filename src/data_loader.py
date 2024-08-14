import os
from io import BytesIO
from typing import BinaryIO
from docx import Document
import fitz  # PyMuPDF

def read_file(file: BytesIO, filename: str) -> str:
    """
    Read content from a file stream based on its extension.

    Args:
        file (BytesIO): The uploaded file object.
        filename (str): The name of the uploaded file.

    Returns:
        str: The text content of the file.

    Raises:
        ValueError: If the file type is unsupported.
    """
    # Get the file extension from the filename
    _, file_extension = os.path.splitext(filename)

    if file_extension.lower() == '.txt':
        return read_txt(file)

    elif file_extension.lower() == '.docx':
        return read_docx(file)

    elif file_extension.lower() == '.pdf':
        return read_pdf(file)

    else:
        raise ValueError(f"Unsupported file extension: {file_extension}")

def read_txt(file: BinaryIO) -> str:
    file.seek(0)
    return file.read().decode('utf-8')

def read_docx(file: BinaryIO) -> str:
    file.seek(0)
    document = Document(file)
    full_text = [paragraph.text for paragraph in document.paragraphs]
    return '\n'.join(full_text)

def read_pdf(file: BinaryIO) -> str:
    file.seek(0)
    document = fitz.open(stream=file.read(), filetype='pdf')
    all_text = [page.get_text() for page in document]
    document.close()
    return '\n'.join(all_text)



# import os
# from io import BytesIO

# import fitz # PyMuPDF
# from docx import Document
# from fastapi import UploadFile

# from typing import BinaryIO


# def read_file(upload_file: UploadFile) -> str:
#     """
#     Read content from a file stream based on its extension.

#     Args:
#         upload_file (UploadFile): The uploaded file object.

#     Returns:
#         str: The text content of the file.

#     Raises:
#         ValueError: If the file type is unsupported.
#     """
#     # Get the file extension from the UploadFile object
#     _, file_extension = os.path.splitext(upload_file.filename)

#     # Convert the SpooledTemporaryFile to BytesIO if needed
#     file_content = BytesIO(upload_file.file.read())
#     upload_file.file.seek(0)  # Reset file pointer after reading

#     if file_extension.lower() == '.txt':
#         return read_txt(file_content)

#     elif file_extension.lower() == '.docx':
#         return read_docx(file_content)

#     elif file_extension.lower() == '.pdf':
#         return read_pdf(file_content)

#     else:
#         raise ValueError(f"Unsupported file extension: {file_extension}")

# def read_txt(file: BinaryIO) -> str:
#     file.seek(0)
#     return file.read().decode('utf-8')

# def read_docx(file: BinaryIO) -> str:
#     file.seek(0)
#     document = Document(file)
#     full_text = [paragraph.text for paragraph in document.paragraphs]
#     return '\n'.join(full_text)

# def read_pdf(file: BinaryIO) -> str:
#     file.seek(0)
#     document = fitz.open(stream=file.read(), filetype='pdf')
#     all_text = [page.get_text() for page in document]
#     document.close()
#     return '\n'.join(all_text)
